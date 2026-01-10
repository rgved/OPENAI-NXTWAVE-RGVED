# main.py
from fastapi import FastAPI, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from backend.grader import grade_answer, companion_feedback
import docx
import fitz
from io import BytesIO
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from io import BytesIO
from PIL import Image
# import pytesseract
import cv2
import numpy as np

# ---------------------------
# QUESTION–ANSWER PARSER (FROM OLD STREAMLIT)
# ---------------------------
import re

def detect_question(line: str):
    keywords = [
        r'\bdefine\b', r'\bdescribe\b', r'\bshow\b', r'\billustrate\b',
        r'\belaborate\b', r'\bexplain\b', r'\bgive\b',
        r'\bwho\b', r'\bwhat\b', r'\bwhere\b',
        r'\bwhen\b', r'\bwhy\b', r'\bhow\b'
    ]
    pattern = r'(\?$|:\s*$|' + "|".join(keywords) + r')'
    return re.search(pattern, line.strip(), flags=re.IGNORECASE)


def smart_parse_text_to_qa(raw_text: str):
    """
    Handles formats like:
    1) Q.1. Question
       ANS: Answer

    2) Question
       Answer paragraph...
    """

    raw_text = re.sub(r'\n+', '\n', raw_text.strip())
    lines = [l.strip() for l in raw_text.split("\n") if l.strip()]

    results = []
    current_q = None
    current_a = []

    for line in lines:
        # ---------- EXPLICIT QUESTION ----------
        q_match = re.match(r'^(Q\.?\s*\d+\.?\s*)(.*)', line, re.IGNORECASE)
        if q_match:
            if current_q:
                results.append({
                    "question": current_q,
                    "student_answer": " ".join(current_a)
                })
            current_q = q_match.group(2)
            current_a = []
            continue

        # ---------- ANSWER PREFIX ----------
        a_match = re.match(r'^(ANS|ANSWER)\s*[:\-]?\s*(.*)', line, re.IGNORECASE)
        if a_match:
            current_a.append(a_match.group(2))
            continue

        # ---------- CONTINUATION ----------
        if current_q:
            current_a.append(line)

    # ---------- FINAL APPEND ----------
    if current_q:
        results.append({
            "question": current_q,
            "student_answer": " ".join(current_a)
        })

    # ---------- 🔥 FALLBACK HEURISTIC ----------
    if not results and len(lines) >= 2:
        return [{
            "question": lines[0],
            "student_answer": " ".join(lines[1:])
        }]

    return results






app = FastAPI(title="Xaminai Backend")

# ---------------------------
# CORS
# ---------------------------
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # tighten later
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---------------------------
# Schemas
# ---------------------------
class GradeRequest(BaseModel):
    question: str
    student_answer: str
    correct_answer: str | None = None
    max_score: int = 5
    difficulty: str = "medium"


class CompanionRequest(BaseModel):
    question: str
    student_answer: str
    correct_answer: str | None = None
    max_score: int = 5


# ---------------------------
# Routes
# ---------------------------
@app.post("/grade")
def grade(req: GradeRequest):
    graded = grade_answer(
        question=req.question,
        student_answer=req.student_answer,
        correct_answer=req.correct_answer or "",
        max_score=req.max_score,
        difficulty=req.difficulty
    )

    # 🔥 SAVE TO DATABASE
    db = SessionLocal()
    record = GradingHistory(
        question=req.question,
        student_answer=req.student_answer,
        score=graded["score"],
        max_score=req.max_score,
        difficulty=req.difficulty,
        mode="text",
        feedback=graded["feedback"]
    )
    db.add(record)
    db.commit()
    db.close()

    return {
         "question": req.question,
        "student_answer": req.student_answer,
        "score": graded["score"],
        "feedback": graded["feedback"]
    }

@app.post("/grade/file")
async def grade_file(
    file: UploadFile = File(...),
    max_score: int = 5,
    difficulty: str = "medium"
):
    raw = await file.read()
    filename = file.filename.lower()

    extracted_text = ""

    try:
        if filename.endswith(".pdf"):
            pdf = fitz.open(stream=raw, filetype="pdf")
            extracted_text = "\n".join(page.get_text() for page in pdf)

        elif filename.endswith(".docx"):
            doc = docx.Document(BytesIO(raw))
            extracted_text = "\n".join(
                p.text.strip() for p in doc.paragraphs if p.text.strip()
            )

        elif filename.endswith(".txt"):
            extracted_text = raw.decode("utf-8", errors="ignore")

        elif filename.endswith(".json"):
            extracted_text = raw.decode("utf-8", errors="ignore")

        else:
            return {
                "score": 0,
                "feedback": f"Unsupported file type: {file.filename}"
            }

    except Exception as e:
        return {
            "score": 0,
            "feedback": f"Failed to read file: {str(e)}"
        }

    if not extracted_text.strip():
        return {
            "score": 0,
            "feedback": "No readable text found in the file."
        }

    # -----------------------------
    # PARSE QUESTION & ANSWER
    # -----------------------------
    parsed = smart_parse_text_to_qa(extracted_text)

    db = SessionLocal()  # 🔥 OPEN DB SESSION ONCE

    # -----------------------------
    # FALLBACK: NO QUESTION FOUND
    # -----------------------------
    if not parsed:
        graded = grade_answer(
            question="Evaluate the following answer",
            student_answer=extracted_text,
            correct_answer="",
            max_score=max_score,
            difficulty=difficulty
        )

        # 🔥 SAVE TO DATABASE
        record = GradingHistory(
            question="Evaluate the following answer",
            student_answer=extracted_text,
            score=graded["score"],
            max_score=max_score,
            difficulty=difficulty,
            mode="file",
            feedback=graded["feedback"]
        )

        db.add(record)
        db.commit()
        db.close()

        return {
            "question": "Evaluate the following answer",
            "student_answer": extracted_text,
            "score": graded["score"],
            "feedback": graded["feedback"]
        }

    # -----------------------------
    # GRADE FIRST QUESTION
    # -----------------------------
    item = parsed[0]

    graded = grade_answer(
        question=item["question"],
        student_answer=item["student_answer"],
        correct_answer="",
        max_score=max_score,
        difficulty=difficulty
    )

    # 🔥 SAVE TO DATABASE
    record = GradingHistory(
        question=item["question"],
        student_answer=item["student_answer"],
        score=graded["score"],
        max_score=max_score,
        difficulty=difficulty,
        mode="file",
        feedback=graded["feedback"]
    )

    db.add(record)
    db.commit()
    db.close()

    return {
        "question": item["question"],
        "student_answer": item["student_answer"],
        "score": graded["score"],
        "feedback": graded["feedback"]
    }

@app.post("/companion")
def companion(req: CompanionRequest):
    return companion_feedback(
        question=req.question,
        student_answer=req.student_answer,
        correct_answer=req.correct_answer or ""
        # max_score=req.max_score
    )

@app.post("/companion/file")
async def companion_file(file: UploadFile = File(...)):
    raw = await file.read()
    filename = file.filename.lower()

    extracted_text = ""

    try:
        # -------- DOCX --------
        if filename.endswith(".docx"):
            doc = docx.Document(BytesIO(raw))
            extracted_text = "\n".join(
                p.text.strip() for p in doc.paragraphs if p.text.strip()
            )

        # -------- PDF --------
        elif filename.endswith(".pdf"):
            pdf = fitz.open(stream=raw, filetype="pdf")
            extracted_text = "\n".join(
                page.get_text().strip() for page in pdf if page.get_text().strip()
            )

        # -------- TXT --------
        elif filename.endswith(".txt"):
            extracted_text = raw.decode("utf-8", errors="ignore")

        # -------- JSON --------
        elif filename.endswith(".json"):
            extracted_text = raw.decode("utf-8", errors="ignore")

        else:
            return {
                "feedback": f"Unsupported file type: {file.filename}",
                "keywords": [],
                "improvement_steps": []
            }

    except Exception as e:
        return {
            "feedback": f"Failed to read file: {str(e)}",
            "keywords": [],
            "improvement_steps": []
        }

    if not extracted_text.strip():
        return {
            "feedback": "No readable text found in the file.",
            "keywords": [],
            "improvement_steps": []
        }

    # 🔥 REAL GEMINI CALL
    return companion_feedback(
        question="Answer extracted from uploaded file",
        student_answer=extracted_text,
        correct_answer=""
    )

# ---------------------------
# GOOGLE DRIVE SETUP
# ---------------------------
SCOPES = ["https://www.googleapis.com/auth/drive.readonly"]
SERVICE_ACCOUNT_FILE = "service_account.json"

import os
import json
from google.oauth2 import service_account

service_account_info = json.loads(
    os.environ["GOOGLE_SERVICE_ACCOUNT_JSON"]
)

credentials = service_account.Credentials.from_service_account_info(
    service_account_info,
    scopes=SCOPES
)


drive_service = build("drive", "v3", credentials=credentials)

def download_drive_file(file_id: str):
    """
    Downloads a file from Google Drive and returns (filename, mime_type, bytes)
    """
    file_meta = drive_service.files().get(
        fileId=file_id,
        fields="name, mimeType"
    ).execute()

    filename = file_meta["name"]
    mime_type = file_meta["mimeType"]

    buffer = BytesIO()

    # Google Docs / Sheets / Slides → export
    if mime_type.startswith("application/vnd.google-apps"):
        # Export Google Docs as DOCX
        request = drive_service.files().export_media(
            fileId=file_id,
            mimeType="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        filename += ".docx"
    else:
        request = drive_service.files().get_media(fileId=file_id)

    downloader = MediaIoBaseDownload(buffer, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()

    buffer.seek(0)
    return filename, mime_type, buffer.read()

from fastapi import Query
from backend.grader import grade_answer


from fastapi import Query

@app.post("/grade/drive")
def grade_from_drive(
    file_id: str = Query(...),
    max_score: int = 5,
    difficulty: str = "medium"
):
    filename, mime_type, raw = download_drive_file(file_id)
    name = filename.lower()

    extracted_text = ""

    try:
        if mime_type == "application/pdf" or name.endswith(".pdf"):
            doc = fitz.open(stream=raw, filetype="pdf")
            extracted_text = "\n".join(page.get_text() for page in doc)

        elif (
            mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            or name.endswith(".docx")
        ):
            d = docx.Document(BytesIO(raw))
            extracted_text = "\n".join(
                p.text.strip() for p in d.paragraphs if p.text.strip()
            )

        elif mime_type == "text/plain" or name.endswith(".txt"):
            extracted_text = raw.decode("utf-8", errors="ignore")

        elif mime_type == "application/json" or name.endswith(".json"):
            extracted_text = raw.decode("utf-8", errors="ignore")

        else:
            return {
                "question": "—",
                "student_answer": "—",
                "score": 0,
                "feedback": f"Unsupported Drive file type: {mime_type}"
            }

    except Exception as e:
        return {
            "question": "—",
            "student_answer": "—",
            "score": 0,
            "feedback": f"Failed to read Drive file: {str(e)}"
        }

    if not extracted_text.strip():
        return {
            "question": "—",
            "student_answer": "—",
            "score": 0,
            "feedback": "No readable text found in Drive file."
        }

    # -----------------------------
    # PARSE QUESTION & ANSWER
    # -----------------------------
    parsed = smart_parse_text_to_qa(extracted_text)

    db = SessionLocal()

    # -----------------------------
    # FALLBACK: NO QUESTION FOUND
    # -----------------------------
    if not parsed:
        graded = grade_answer(
            question="Evaluate the following answer",
            student_answer=extracted_text,
            correct_answer="",
            max_score=max_score,
            difficulty=difficulty
        )

        record = GradingHistory(
            question="Evaluate the following answer",
            student_answer=extracted_text,
            score=graded["score"],
            max_score=max_score,
            difficulty=difficulty,
            mode="drive",
            feedback=graded["feedback"]
        )

        db.add(record)
        db.commit()
        db.close()

        return {
            "question": "Evaluate the following answer",
            "student_answer": extracted_text,
            "score": graded["score"],
            "feedback": graded["feedback"]
        }

    # -----------------------------
    # GRADE FIRST QUESTION
    # -----------------------------
    item = parsed[0]

    graded = grade_answer(
        question=item["question"],
        student_answer=item["student_answer"],
        correct_answer="",
        max_score=max_score,
        difficulty=difficulty
    )

    record = GradingHistory(
        question=item["question"],
        student_answer=item["student_answer"],
        score=graded["score"],
        max_score=max_score,
        difficulty=difficulty,
        mode="drive",
        feedback=graded["feedback"]
    )

    db.add(record)
    db.commit()
    db.close()

    return {
        "question": item["question"],
        "student_answer": item["student_answer"],
        "score": graded["score"],
        "feedback": graded["feedback"]
    }


from fastapi import UploadFile, File
from backend.grader import grade_from_image


@app.post("/grade/image")
async def grade_image(
    file: UploadFile = File(...),
    max_score: int = 5,
    difficulty: str = "medium"
):
    image_bytes = await file.read()

    graded = grade_from_image(
        image_bytes=image_bytes,
        max_score=max_score,
        difficulty=difficulty
    )

    # 🔥 SAVE TO DATABASE
    db = SessionLocal()
    record = GradingHistory(
        question="Answer extracted from image",
        student_answer="(Image upload)",
        score=graded["score"],
        max_score=max_score,
        difficulty=difficulty,
        mode="image",
        feedback=graded["feedback"]
    )
    db.add(record)
    db.commit()
    db.close()

    return graded

from fastapi import UploadFile, File
from backend.grader import companion_from_image

@app.post("/companion/image")
async def companion_image(file: UploadFile = File(...)):
    image_bytes = await file.read()

    return companion_from_image(image_bytes)



from fastapi import UploadFile, File
from google.oauth2 import service_account
from googleapiclient.discovery import build
from io import BytesIO
import docx
import json
import fitz  # PyMuPDF

SCOPES = ["https://www.googleapis.com/auth/drive.readonly"]
SERVICE_ACCOUNT_FILE = "service_account.json"


def get_drive_service():
    creds = service_account.Credentials.from_service_account_file(
        SERVICE_ACCOUNT_FILE,
        scopes=SCOPES
    )
    return build("drive", "v3", credentials=creds)


@app.post("/companion/drive")
def companion_from_drive(file_id: str):
    service = get_drive_service()

    file_meta = service.files().get(
        fileId=file_id,
        fields="name, mimeType"
    ).execute()

    name = file_meta["name"].lower()
    mime = file_meta["mimeType"]

    request = (
        service.files().export_media(fileId=file_id, mimeType="application/pdf")
        if mime.startswith("application/vnd.google-apps")
        else service.files().get_media(fileId=file_id)
    )

    fh = BytesIO()
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()

    fh.seek(0)
    text = ""

    try:
        # ---------- PDF ----------
        if name.endswith(".pdf"):
            doc = fitz.open(stream=fh.read(), filetype="pdf")
            text = "\n".join(page.get_text() for page in doc)

        # ---------- DOCX ----------
        elif name.endswith(".docx"):
            doc = docx.Document(fh)
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

        # ---------- TXT / JSON ----------
        elif name.endswith(".txt") or name.endswith(".json"):
            text = fh.read().decode("utf-8", errors="ignore")

        else:
            return {
                "feedback": "Unsupported file type from Drive",
                "keywords": [],
                "improvement_steps": []
            }

    except Exception as e:
        return {
            "feedback": f"Failed to read Drive file: {str(e)}",
            "keywords": [],
            "improvement_steps": []
        }

    if not text.strip():
        return {
            "feedback": "No readable text found in Drive file",
            "keywords": [],
            "improvement_steps": []
        }

    return companion_feedback(
        question="Answer extracted from Google Drive file",
        student_answer=text,
        correct_answer=""
    )

#DASHBOARD DATABASE

from fastapi import Query
from backend.database import SessionLocal
from backend.models import GradingHistory

from fastapi import Query
from backend.database import SessionLocal
from backend.models import GradingHistory



@app.get("/dashboard/history")
def get_grading_history(
    mode: str | None = Query(default=None)
):
    db = SessionLocal()

    query = db.query(GradingHistory)

    if mode:
        query = query.filter(GradingHistory.mode == mode)

    records = query.order_by(
        GradingHistory.created_at.desc()
    ).all()

    db.close()

    return [
        {
            "id": r.id,
            "question": r.question,
            "student_answer": r.student_answer,
            "score": r.score,
            "max_score": r.max_score,
            "difficulty": r.difficulty,
            "mode": r.mode,
            "feedback": r.feedback,
            "created_at": r.created_at.isoformat()
        }
        for r in records
    ]



import csv
from fastapi.responses import StreamingResponse
from io import StringIO

@app.get("/dashboard/export")
def export_history(mode: str | None = Query(default=None)):
    db = SessionLocal()

    query = db.query(GradingHistory)
    if mode:
        query = query.filter(GradingHistory.mode == mode)

    records = query.order_by(GradingHistory.created_at.desc()).all()
    db.close()

    buffer = StringIO()
    writer = csv.writer(buffer)

    # Header
    writer.writerow([
        "Question",
        "Student Answer",
        "Score",
        "Max Score",
        "Difficulty",
        "Mode",
        "Feedback",
        "Created At"
    ])

    for r in records:
        writer.writerow([
            r.question,
            r.student_answer,
            r.score,
            r.max_score,
            r.difficulty,
            r.mode,
            r.feedback,
            r.created_at.isoformat()
        ])

    buffer.seek(0)

    filename = f"grading_history_{mode or 'all'}.csv"

    return StreamingResponse(
        buffer,
        media_type="text/csv",
        headers={
            "Content-Disposition": f"attachment; filename={filename}"
        }
    )
